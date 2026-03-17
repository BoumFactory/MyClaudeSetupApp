# -*- coding: utf-8 -*-
"""
generate_docx.py : Génération de documents DOCX style F. Loiseau depuis JSON
═══════════════════════════════════════════════════════════════════════════════

UTILISATION
───────────
    python generate_docx.py chemin/vers/config.json [--output chemin/vers/sortie]

DESCRIPTION
───────────
Ce script lit un fichier JSON contenant la description structurée d'un document
pédagogique au style F. Loiseau et génère :
  - Version élève (DOCX)
  - Version corrigée (DOCX)
  - PDFs correspondants (via LibreOffice ou Word COM)

Le JSON doit respecter le schéma défini dans:
  .claude/skills/docx-loiseau/references/json_schema.md

EXEMPLE DE JSON
───────────────
{
  "metadata": {
    "doc_type": "Fiche Méthode",
    "number": "01",
    "title": "Dériver une fonction",
    "level": "1ère",
    "author": "M. Dupont",
    "doc_mode": "exercice"
  },
  "versions": ["eleve", "corrige"],
  "pages": [
    {
      "components": [
        {"type": "header"},
        {"type": "section_title", "number": 1, "title": "Rappels"},
        {"type": "content_box", "label": "Règle", "content": "..."},
        {"type": "exercise", "number": 1, "statement": "...", "items": [...]}
      ]
    }
  ]
}

COMPOSANTS DISPONIBLES
──────────────────────
- header             : En-tête Loiseau (type, numéro, titre, badge niveau)
- section_title      : Titre de section numéroté
- section_band       : Bandeau dégradé noir→gris
- content_box        : Boîte avec label vertical noir
- exercise           : Exercice avec sous-questions et grille réponse
- items              : Liste de questions
- answer_grid        : Grille Séyès pour réponses élèves
- table              : Tableau standard
- text               : Paragraphe texte
- page_break         : Saut de page
- spacer             : Espace vertical

CONFIGURATION SKILL
───────────────────
Chemin grille Séyès   : SKILL_DIR / "assets" / "grille_seyes.png"
Chemin MML2OMML.XSL   : C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL

NOTES D'UTILISATION
───────────────────
1. Toutes les formules LaTeX sont entre $ ... $ et converties en OMML
2. Les mots-clés entre ** ... ** sont en Bold+Italic
3. La version élève masque les corrections et affiche des grilles Séyès
4. La version corrigée affiche les corrections en rouge
5. Le script génère automatiquement 2 fichiers DOCX (élève + corrigé)
6. L'export PDF est automatique via LibreOffice, fallback sur Word COM

ACCENTS FRANÇAIS
────────────────
IMPORTANT : Tous les accents français (é, è, ê, à, ù, ç, etc.) doivent être
correctement encodés en UTF-8. Le script utilise sys.stdout.reconfigure(encoding='utf-8').

═══════════════════════════════════════════════════════════════════════════════
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

import os
import re
import json
import argparse
import subprocess
from pathlib import Path
from copy import deepcopy
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import latex2mathml.converter
from PIL import Image, ImageDraw, ImageFont
import tempfile


# ════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════

SKILL_DIR = Path(__file__).parent.parent
ASSETS_DIR = SKILL_DIR / "assets"
GRILLE_SEYES_PNG = ASSETS_DIR / "grille_seyes.png"
MML2OMML_XSL = Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL")

# Couleurs Loiseau
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY_LIGHT = "F2F2F2"
GREY_MEDIUM = "D9D9D9"
RED = RGBColor(0xFF, 0x00, 0x00)
GREEN = RGBColor(0x00, 0xB0, 0x50)
VIOLET = "7030A0"
VIOLET_LIGHT = "D3B5E9"

# Polices
FONT_COOPER = "Cooper"
FONT_COOPER_BLACK = "Cooper Black"
FONT_CAMBRIA = "Cambria"
FONT_DEFAULT = "Calibri"

# Chemin vers la police Cooper Black pour Pillow
COOPER_BLACK_TTF = Path(os.environ.get('WINDIR', r'C:\Windows')) / 'Fonts' / 'COOPBL.TTF'


# ════════════════════════════════════════════════════════════════════════════
# BADGE NIVEAU (carré arrondi dégradé violet, style Loiseau)
# ════════════════════════════════════════════════════════════════════════════

def create_level_badge(level="4e", size=400):
    """
    Génère un badge carré arrondi avec dégradé violet et texte blanc,
    reproduisant fidèlement le badge original F. Loiseau.

    Args:
        level: String du niveau (ex: "1ère", "4e", "3e")
        size: Taille du badge en pixels (défaut 400)

    Returns:
        Chemin vers le fichier temporaire PNG
    """
    img = Image.new('RGBA', (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    margin = int(size * 0.05)
    radius = int(size * 0.18)

    draw.rounded_rectangle(
        [margin, margin, size - margin, size - margin],
        radius=radius,
        fill=(80, 60, 130),
        outline=(45, 27, 105),
        width=int(size * 0.04)
    )

    inner_m = margin + int(size * 0.06)
    inner_r = int(radius * 0.7)
    for i in range(20):
        t = i / 19.0
        r = int(229 - t * 70)
        g = int(213 - t * 70)
        b = int(255 - t * 50)
        shrink = int(t * size * 0.15)
        x0 = inner_m + shrink
        y0 = inner_m + shrink
        x1 = size - inner_m - shrink
        y1 = size - inner_m - shrink
        if x1 > x0 and y1 > y0:
            draw.rounded_rectangle(
                [x0, y0, x1, y1],
                radius=max(1, inner_r - int(t * 10)),
                fill=(r, g, b)
            )

    match = re.match(r'^(\d+)(.*)', level)
    if match:
        main_text = match.group(1)
        sup_text = match.group(2)
    else:
        main_text = level
        sup_text = ""

    try:
        font_main = ImageFont.truetype(str(COOPER_BLACK_TTF), int(size * 0.55))
        font_sup = ImageFont.truetype(str(COOPER_BLACK_TTF), int(size * 0.25))
    except Exception:
        font_main = ImageFont.load_default()
        font_sup = ImageFont.load_default()

    bbox_main = draw.textbbox((0, 0), main_text, font=font_main)
    main_w = bbox_main[2] - bbox_main[0]
    main_h = bbox_main[3] - bbox_main[1]
    bbox_sup = draw.textbbox((0, 0), sup_text, font=font_sup) if sup_text else (0, 0, 0, 0)
    sup_w = bbox_sup[2] - bbox_sup[0]

    total_w = main_w + sup_w + int(size * 0.01)
    x_start = (size - total_w) // 2
    y_main = (size - main_h) // 2 - int(size * 0.04)

    draw.text((x_start + 2, y_main + 2), main_text, fill=(60, 40, 100, 120), font=font_main)
    draw.text((x_start, y_main), main_text, fill=(255, 255, 255), font=font_main)

    if sup_text:
        x_sup = x_start + main_w + int(size * 0.01)
        y_sup = y_main + int(size * 0.08)
        draw.text((x_sup + 1, y_sup + 1), sup_text, fill=(60, 40, 100, 120), font=font_sup)
        draw.text((x_sup, y_sup), sup_text, fill=(255, 255, 255), font=font_sup)

    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='badge_')
    img.save(tmp.name, 'PNG')
    tmp.close()
    return tmp.name


# ════════════════════════════════════════════════════════════════════════════
# HELPERS OMML (LaTeX → Word Math)
# ════════════════════════════════════════════════════════════════════════════

_xslt_transform = None

def get_xslt_transform():
    """Charger et cacher le XSLT de conversion MathML → OMML"""
    global _xslt_transform
    if _xslt_transform is None:
        if MML2OMML_XSL.exists():
            xslt_tree = etree.parse(str(MML2OMML_XSL))
            _xslt_transform = etree.XSLT(xslt_tree)
        else:
            raise FileNotFoundError(f"MML2OMML.XSL introuvable: {MML2OMML_XSL}")
    return _xslt_transform


def latex_to_omml(latex_str):
    """Convertir une formule LaTeX en OMML pour Word"""
    mathml_str = latex2mathml.converter.convert(latex_str)
    transform = get_xslt_transform()
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = transform(mathml_tree)

    omml_root = omml_tree.getroot()
    omath = omml_root
    if omml_root.tag.endswith('oMathPara'):
        omath_list = omml_root.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
        if omath_list:
            omath = omath_list[0]
    elif not omml_root.tag.endswith('oMath'):
        omath_list = omml_root.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
        if omath_list:
            omath = omath_list[0]

    return omath


def add_math_to_paragraph(paragraph, latex_str):
    """
    Ajouter une formule LaTeX (convertie en OMML) à un paragraphe existant.
    Gère les accents français dans le texte mathématique et fournit un
    fallback lisible en cas d'échec de conversion.
    """
    # Nettoyer la formule : espaces superflus
    latex_clean = latex_str.strip()
    if not latex_clean:
        return

    try:
        omml = latex_to_omml(latex_clean)
        paragraph._element.append(omml)
    except Exception as e:
        # Fallback : afficher la formule en italique entre crochets
        # Tenter de rendre lisible (remplacer les commandes LaTeX courantes)
        fallback = latex_clean
        fallback = fallback.replace('\\times', '×')
        fallback = fallback.replace('\\cdot', '·')
        fallback = fallback.replace('\\leq', '≤')
        fallback = fallback.replace('\\geq', '≥')
        fallback = fallback.replace('\\neq', '≠')
        fallback = fallback.replace('\\pm', '±')
        fallback = fallback.replace('\\infty', '∞')
        fallback = fallback.replace('\\mathbb{R}', 'ℝ')
        fallback = fallback.replace('\\mathbb{N}', 'ℕ')
        fallback = fallback.replace('\\mathbb{Z}', 'ℤ')
        run = paragraph.add_run(f" {fallback} ")
        run.font.italic = True
        run.font.size = Pt(10)
        print(f"  [WARN] Formule non convertie: {latex_clean[:40]}... ({e})")


def add_text_with_math(paragraph, text, font_name=None, font_size=None, bold=None, italic=None, color=None):
    """
    Ajouter du texte mélangé avec des formules LaTeX.
    Les formules sont entre $ ... $ (inline).
    Les mots entre ** ... ** sont en Bold+Italic (convention Loiseau).
    """
    parts = re.split(r'(\$[^$]+\$|\*\*[^*]+\*\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            latex = part[1:-1]
            add_math_to_paragraph(paragraph, latex)
        elif part.startswith('**') and part.endswith('**'):
            keyword = part[2:-2]
            run = paragraph.add_run(keyword)
            run.bold = True
            run.italic = True
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if color:
                run.font.color.rgb = color
        else:
            run = paragraph.add_run(part)
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if color:
                run.font.color.rgb = color


# ════════════════════════════════════════════════════════════════════════════
# HELPERS FORMATAGE DOCX
# ════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Appliquer un fond à une cellule"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear" w:color="auto"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_vertical_alignment(cell, align="center"):
    """Alignement vertical d'une cellule"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)


def set_cell_width(cell, width_twips):
    """Définir la largeur d'une cellule en twips"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    existing = tcPr.findall(qn('w:tcW'))
    for e in existing:
        tcPr.remove(e)
    tcPr.append(tcW)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Définir les bordures d'une cellule"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val == 'nil':
            borders_xml += f'<w:{side} w:val="nil"/>'
        elif val == 'none':
            borders_xml += f'<w:{side} w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    borders_xml += '</w:tcBorders>'

    borders = parse_xml(borders_xml)
    existing = tcPr.findall(qn('w:tcBorders'))
    for e in existing:
        tcPr.remove(e)
    tcPr.append(borders)


def remove_table_borders(table):
    """Supprimer toutes les bordures d'une table"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>'''
    borders = parse_xml(borders_xml)

    existing = tblPr.findall(qn('w:tblBorders'))
    for e in existing:
        tblPr.remove(e)
    tblPr.append(borders)


def set_table_width_pct(table, pct=5000):
    """Définir la largeur de table en pourcentage (5000 = 100%)"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{pct}" w:type="pct"/>')
    existing = tblPr.findall(qn('w:tblW'))
    for e in existing:
        tblPr.remove(e)
    tblPr.append(tblW)


def set_cell_text_direction(cell, direction="btLr"):
    """Définir la direction du texte dans une cellule (btLr = rotation 90°)"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.findall(qn('w:textDirection'))
    for e in existing:
        tcPr.remove(e)
    td = parse_xml(f'<w:textDirection {nsdecls("w")} w:val="{direction}"/>')
    tcPr.append(td)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Définir l'espacement d'un paragraphe"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def add_page_break(doc):
    """Ajouter un saut de page"""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)
    set_paragraph_spacing(para, 0, 0)
    return para


# ════════════════════════════════════════════════════════════════════════════
# COMPOSANTS LOISEAU
# ════════════════════════════════════════════════════════════════════════════

def add_header_table(doc, doc_type, number, title, level="1ère"):
    """
    Ajouter l'en-tête Loiseau : table 1×3
    [Type + N°] [Titre] [Badge niveau]
    """
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    set_table_width_pct(table, 5000)

    cell0 = table.cell(0, 0)
    set_cell_shading(cell0, GREY_MEDIUM)
    set_cell_vertical_alignment(cell0, "center")
    set_cell_width(cell0, 1241)

    p1 = cell0.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p1, 0, 0)
    doc_type_clean = doc_type.replace('\n', ' ')
    run1 = p1.add_run(doc_type_clean)
    run1.font.name = FONT_COOPER
    run1.font.size = Pt(10)

    p2 = cell0.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, 0, 0)
    run2 = p2.add_run(str(number).zfill(2))
    run2.font.name = FONT_COOPER
    run2.font.size = Pt(11)

    cell1 = table.cell(0, 1)
    set_cell_shading(cell1, GREY_LIGHT)
    set_cell_vertical_alignment(cell1, "center")
    set_cell_width(cell1, 8443)
    set_cell_borders(cell1, left='nil')

    p_title = cell1.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_title, 2, 2)
    run_title = p_title.add_run(title)
    run_title.font.name = FONT_COOPER_BLACK
    run_title.font.size = Pt(18)

    cell2 = table.cell(0, 2)
    set_cell_vertical_alignment(cell2, "center")
    set_cell_width(cell2, 1088)

    p_badge = cell2.paragraphs[0]
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_badge, 0, 0)

    badge_path = create_level_badge(level=level, size=400)
    run_badge = p_badge.add_run()
    run_badge.add_picture(badge_path, width=Cm(1.3))
    try:
        os.unlink(badge_path)
    except Exception:
        pass

    return table


def add_section_title(doc, number, title):
    """Ajouter un titre de section numérotée"""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell0 = table.cell(0, 0)
    set_cell_shading(cell0, "000000")
    set_cell_vertical_alignment(cell0, "center")
    set_cell_width(cell0, 562)

    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p0, 2, 2)
    run0 = p0.add_run(str(number))
    run0.font.name = FONT_CAMBRIA
    run0.font.size = Pt(16)
    run0.bold = True
    run0.font.color.rgb = WHITE

    cell1 = table.cell(0, 1)
    set_cell_shading(cell1, GREY_LIGHT)
    set_cell_vertical_alignment(cell1, "center")
    set_cell_width(cell1, 10200)

    p1 = cell1.paragraphs[0]
    set_paragraph_spacing(p1, 2, 2)
    run1 = p1.add_run(title)
    run1.font.name = FONT_CAMBRIA
    run1.font.size = Pt(16)
    run1.bold = True

    return table


def add_content_box(doc, label, content_text, subtitle=""):
    """
    Ajouter une boîte de contenu style Loiseau :
    [Label vertical noir à gauche] [Contenu sur fond gris à droite]
    """
    table = doc.add_table(rows=1, cols=2)

    cell_label = table.cell(0, 0)
    set_cell_shading(cell_label, "000000")
    set_cell_width(cell_label, 480)
    set_cell_vertical_alignment(cell_label, "center")
    set_cell_text_direction(cell_label, "btLr")

    p_label = cell_label.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_label, 0, 0)
    run_label = p_label.add_run(label)
    run_label.font.name = FONT_CAMBRIA
    run_label.font.size = Pt(11)
    run_label.bold = True
    run_label.font.color.rgb = WHITE

    cell_content = table.cell(0, 1)
    set_cell_shading(cell_content, GREY_LIGHT)
    set_cell_width(cell_content, 10282)

    if subtitle:
        p_sub = cell_content.paragraphs[0]
        set_paragraph_spacing(p_sub, 2, 2)
        pPr_sub = p_sub._element.get_or_add_pPr()
        shd_sub = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{GREY_MEDIUM}" w:val="clear" w:color="auto"/>')
        pPr_sub.append(shd_sub)
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = FONT_CAMBRIA
        run_sub.font.size = Pt(11)
        run_sub.bold = True
        run_sub.italic = True
        p_content = cell_content.add_paragraph()
    else:
        p_content = cell_content.paragraphs[0]

    set_paragraph_spacing(p_content, 4, 4)
    add_text_with_math(p_content, content_text)

    return table


def create_section_band_image(title, width_cm=19.0, height_cm=0.7):
    """
    Crée une image de bandeau de section style Loiseau :
    Rectangle arrondi avec dégradé horizontal noir → gris → quasi blanc.
    """
    dpi = 150
    px_per_cm = dpi / 2.54
    w = int(width_cm * px_per_cm)
    h = int(height_cm * px_per_cm)

    img = Image.new('RGBA', (w, h), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    radius = int(h * 0.4)

    for x in range(w):
        t = x / (w - 1)
        if t < 0.6:
            tt = t / 0.6
            r = int(0 + tt * 214)
            g = int(0 + tt * 214)
            b = int(0 + tt * 214)
        else:
            tt = (t - 0.6) / 0.4
            r = int(214 + tt * 28)
            g = int(214 + tt * 28)
            b = int(214 + tt * 28)
        draw.line([(x, 0), (x, h - 1)], fill=(r, g, b))

    mask = Image.new('L', (w, h), 0)
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle([0, 0, w - 1, h - 1], radius=radius, fill=255)
    img.putalpha(mask)

    try:
        font = ImageFont.truetype(
            str(Path(os.environ.get('WINDIR', r'C:\Windows')) / 'Fonts' / 'ARLRDBD.TTF'),
            int(h * 0.55)
        )
    except Exception:
        try:
            font = ImageFont.truetype(str(COOPER_BLACK_TTF), int(h * 0.55))
        except Exception:
            font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), title, font=font)
    text_h = bbox[3] - bbox[1]
    y_text = (h - text_h) // 2 - int(h * 0.05)
    draw.text((int(w * 0.02), y_text), title, fill=(255, 255, 255), font=font)

    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='band_')
    img.save(tmp.name, 'PNG')
    tmp.close()
    return tmp.name


def add_section_band(doc, title):
    """Ajouter une bande de section style Loiseau"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, 6, 1)

    band_path = create_section_band_image(title)
    run = para.add_run()
    run.add_picture(band_path, width=Cm(19.0))
    try:
        os.unlink(band_path)
    except Exception:
        pass

    return para


def add_sub_label_run(paragraph, label_text, is_correction=False):
    """Ajouter un label de sous-question avec petit fond coloré"""
    run = paragraph.add_run(f" {label_text} ")
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = WHITE

    rPr = run._element.get_or_add_rPr()
    fill_color = "FF0000" if is_correction else "000000"
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear" w:color="auto"/>')
    rPr.append(shd)

    run_sp = paragraph.add_run("  ")
    run_sp.font.size = Pt(10)
    return run


def create_seyes_grid(width_cm=19.0, height_cm=5.0):
    """
    Crée une image de grille Séyès aux dimensions demandées par découpage/tuilage
    de l'image source originale.
    """
    if not GRILLE_SEYES_PNG.exists():
        return None

    source = Image.open(str(GRILLE_SEYES_PNG))
    src_w, src_h = source.size

    px_per_cm = src_w / 17.2

    target_w = int(width_cm * px_per_cm)
    target_h = int(height_cm * px_per_cm)

    result = Image.new('RGBA', (target_w, target_h), (255, 255, 255, 0))

    for x_offset in range(0, target_w, src_w):
        for y_offset in range(0, target_h, src_h):
            result.paste(source, (x_offset, y_offset))

    result = result.crop((0, 0, target_w, target_h))

    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False, prefix='grille_')
    result.save(tmp.name, 'PNG')
    tmp.close()
    return tmp.name


def add_answer_grid(doc, height_cm=5.0, width_cm=19.0):
    """Ajouter une grille Séyès pour réponses élèves"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, 1, 1)

    grid_path = create_seyes_grid(width_cm=width_cm, height_cm=height_cm)

    if grid_path:
        run = para.add_run()
        run.add_picture(grid_path, width=Cm(width_cm), height=Cm(height_cm))
        try:
            os.unlink(grid_path)
        except Exception:
            pass
    else:
        print(f"  [WARN] Grille Séyès non trouvée: {GRILLE_SEYES_PNG}")
        table = doc.add_table(rows=int(height_cm * 2), cols=1)
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders_xml = f'''<w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="7F7F7F"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="7F7F7F"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="7F7F7F"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="7F7F7F"/>
            <w:insideH w:val="single" w:sz="2" w:space="0" w:color="AFEEEE"/>
        </w:tblBorders>'''
        existing = tblPr.findall(qn('w:tblBorders'))
        for e in existing:
            tblPr.remove(e)
        tblPr.append(parse_xml(borders_xml))
        set_table_width_pct(table, 5000)
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="284" w:hRule="exact"/>'))
        return table

    return para


def add_exercise_label(doc, number, statement, points=None, difficulty=None, doc_mode="exercice"):
    """
    Ajouter le label exercice style Loiseau.
    Format: Bullet + "Exercice N :" + énoncé + points/étoiles si applicable
    """
    para = doc.add_paragraph()
    set_paragraph_spacing(para, 4, 1)

    run_bullet = para.add_run()
    run_bullet.font.name = FONT_COOPER_BLACK
    run_bullet.font.size = Pt(11)
    sym_xml = f'<w:sym {nsdecls("w")} w:font="Wingdings" w:char="F0DC"/>'
    run_bullet._element.append(parse_xml(sym_xml))

    run_space = para.add_run(" ")
    run_space.font.name = FONT_COOPER_BLACK
    run_space.font.size = Pt(11)

    run_label = para.add_run(f"Exercice {number} ")
    run_label.font.name = FONT_COOPER
    run_label.font.size = Pt(11)

    if points and doc_mode == "eval":
        pPr = para._element.get_or_add_pPr()
        tabs_xml = f'<w:tabs {nsdecls("w")}><w:tab w:val="right" w:leader="underscore" w:pos="10772"/></w:tabs>'
        pPr.append(parse_xml(tabs_xml))

        run_tab = para.add_run()
        run_tab._element.append(parse_xml(f'<w:tab {nsdecls("w")}/>'))

        run_pts = para.add_run(f"{points} point{'s' if float(points.replace(',', '.')) > 1 else ''}")
        run_pts.font.name = "Arial Rounded MT Bold"
        run_pts.font.size = Pt(11)
    else:
        run_colon = para.add_run(": ")
        run_colon.font.name = FONT_COOPER
        run_colon.font.size = Pt(11)

    if difficulty and doc_mode in ("exercice", "cours"):
        stars = "★" * difficulty + "☆" * (3 - difficulty)
        run_stars = para.add_run(f"  {stars}")
        run_stars.font.size = Pt(9)
        run_stars.font.color.rgb = RGBColor(0x70, 0x30, 0xA0)

    if statement:
        add_text_with_math(para, statement)

    return para


def add_small_spacer(doc, height_pt=4):
    """Ajouter un petit espace"""
    para = doc.add_paragraph()
    set_paragraph_spacing(para, 0, 0)
    run = para.add_run()
    run.font.size = Pt(height_pt)
    return para


def add_footer(doc, author, doc_type, number, title, level="1ère"):
    """Configurer le pied de page style Loiseau"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p.clear()

        tbl_element = parse_xml(
            f'<w:tbl {nsdecls("w")}>'
            f'  <w:tblPr>'
            f'    <w:tblW w:w="5000" w:type="pct"/>'
            f'    <w:tblBorders>'
            f'      <w:top w:val="single" w:sz="8" w:space="0" w:color="auto"/>'
            f'    </w:tblBorders>'
            f'  </w:tblPr>'
            f'  <w:tblGrid><w:gridCol w:w="5000"/><w:gridCol w:w="5000"/><w:gridCol w:w="800"/></w:tblGrid>'
            f'  <w:tr>'
            f'    <w:tc><w:p/></w:tc>'
            f'    <w:tc><w:p/></w:tc>'
            f'    <w:tc><w:p/></w:tc>'
            f'  </w:tr>'
            f'</w:tbl>'
        )
        footer._element.append(tbl_element)

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        cells = tbl_element.findall('.//w:tc', ns)

        p_left = cells[0].find('w:p', ns)
        run_author_xml = (
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr><w:i/><w:sz w:val="18"/></w:rPr>'
            f'  <w:t>{author}</w:t>'
            f'</w:r>'
        )
        p_left.append(parse_xml(run_author_xml))

        p_right = cells[1].find('w:p', ns)
        pPr_right = parse_xml(f'<w:pPr {nsdecls("w")}><w:jc w:val="right"/></w:pPr>')
        p_right.insert(0, pPr_right)
        info_text = f"{level} \u2013 {doc_type} {str(number).zfill(2)} \u2013 {title}"
        run_info_xml = (
            f'<w:r {nsdecls("w")}>'
            f'  <w:rPr><w:i/><w:sz w:val="18"/></w:rPr>'
            f'  <w:t xml:space="preserve">{info_text}</w:t>'
            f'</w:r>'
        )
        p_right.append(parse_xml(run_info_xml))

        p_page = cells[2].find('w:p', ns)
        pPr_page = parse_xml(f'<w:pPr {nsdecls("w")}><w:jc w:val="right"/></w:pPr>')
        p_page.insert(0, pPr_page)

        page_field_xml = (
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:fldChar w:fldCharType="end"/></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:t>/</w:t></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:instrText xml:space="preserve"> NUMPAGES </w:instrText></w:r>'
            f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/></w:rPr>'
            f'<w:fldChar w:fldCharType="end"/></w:r>'
        )
        wrapper = f'<w:p {nsdecls("w")}>{page_field_xml}</w:p>'
        parsed = parse_xml(wrapper)
        for child in list(parsed):
            p_page.append(child)


# ════════════════════════════════════════════════════════════════════════════
# RENDU DES COMPOSANTS (DISPATCHER JSON → DOCX)
# ════════════════════════════════════════════════════════════════════════════

def render_component(doc, component, metadata, is_corrige):
    """
    Dispatcher qui transforme un composant JSON en contenu DOCX.

    Args:
        doc: Document Word
        component: Dict du composant JSON
        metadata: Dict des metadata du document
        is_corrige: Boolean (version corrigée ou élève)
    """
    comp_type = component.get("type", "").lower()

    if comp_type == "header":
        add_header_table(
            doc,
            metadata["doc_type"],
            metadata["number"],
            metadata["title"],
            metadata["level"]
        )

    elif comp_type == "spacer":
        height = component.get("height_pt", 4)
        add_small_spacer(doc, height)

    elif comp_type == "section_title":
        add_section_title(
            doc,
            component.get("number", 1),
            component.get("title", "")
        )

    elif comp_type == "section_band":
        add_section_band(doc, component.get("title", ""))

    elif comp_type == "content_box":
        add_content_box(
            doc,
            component.get("label", ""),
            component.get("content", ""),
            component.get("subtitle", "")
        )

    elif comp_type == "text":
        para = doc.add_paragraph()
        set_paragraph_spacing(para, 4, 2)
        add_text_with_math(para, component.get("content", ""))

    elif comp_type == "exercise":
        number = component.get("number", 1)
        statement = component.get("statement", "")
        points = component.get("points")
        difficulty = component.get("difficulty")
        items = component.get("items", [])
        answer_grid_height = component.get("answer_grid_height", 5.0)
        doc_mode = metadata.get("doc_mode", "exercice")

        add_exercise_label(doc, number, statement, points, difficulty, doc_mode)

        if items:
            for item in items:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, 2, 2)
                add_sub_label_run(p, item.get("label", ""), is_correction=False)
                add_text_with_math(p, item.get("content", ""))

                if is_corrige and "correction" in item:
                    p_corr = doc.add_paragraph()
                    set_paragraph_spacing(p_corr, 2, 2)
                    add_sub_label_run(p_corr, item.get("label", ""), is_correction=True)
                    add_text_with_math(p_corr, item["correction"])

        if not is_corrige and answer_grid_height > 0:
            add_answer_grid(doc, height_cm=answer_grid_height)

    elif comp_type == "items":
        items = component.get("items", [])
        layout = component.get("layout", "list")

        if layout == "grid":
            table = doc.add_table(rows=(len(items) + 1) // 2, cols=2)
            set_table_width_pct(table, 5000)
            for idx, item in enumerate(items):
                row = idx // 2
                col = idx % 2
                cell = table.cell(row, col)
                p = cell.paragraphs[0]
                set_paragraph_spacing(p, 4, 4)
                add_sub_label_run(p, item.get("label", ""))
                add_text_with_math(p, item.get("content", ""))
        else:  # list
            for item in items:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, 2, 2)
                add_sub_label_run(p, item.get("label", ""))
                add_text_with_math(p, item.get("content", ""))

    elif comp_type == "answer_grid":
        height = component.get("height_cm", 5.0)
        width = component.get("width_cm", 19.0)
        add_answer_grid(doc, height_cm=height, width_cm=width)

    elif comp_type == "table":
        headers = component.get("headers", [])
        rows = component.get("rows", [])

        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        set_table_width_pct(table, 5000)

        for col_idx, header in enumerate(headers):
            cell = table.cell(0, col_idx)
            set_cell_shading(cell, "000000")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, 4, 4)
            add_text_with_math(p, header, font_name=FONT_CAMBRIA, font_size=Pt(11),
                               bold=True, color=WHITE)

        for row_idx, row_data in enumerate(rows, 1):
            bg = GREY_LIGHT if row_idx % 2 == 1 else "FFFFFF"
            for col_idx, content in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                set_cell_shading(cell, bg)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, 4, 4)
                add_text_with_math(p, content, font_size=Pt(11))

    elif comp_type == "page_break":
        add_page_break(doc)

    else:
        print(f"  [WARN] Type de composant inconnu: {comp_type}")


# ════════════════════════════════════════════════════════════════════════════
# CRÉATION DE DOCUMENT DEPUIS JSON
# ════════════════════════════════════════════════════════════════════════════

def create_document_from_json(data, version="eleve"):
    """
    Crée un document Word à partir d'une structure JSON.

    Args:
        data: Dict contenant la structure JSON du document
        version: "eleve" ou "corrige"

    Returns:
        Document Word configuré
    """
    is_corrige = version == "corrige"
    metadata = data.get("metadata", {})
    pages = data.get("pages", [])

    doc = Document()

    # Configuration page
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    # Style Normal
    style = doc.styles['Normal']
    style.font.name = FONT_DEFAULT
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Types de composants « majeurs » qui nécessitent un spacer automatique entre eux.
    # Le spacer (2pt) est inséré automatiquement ENTRE deux composants majeurs consécutifs.
    # L'agent n'a PAS besoin de mettre des spacers dans le JSON.
    MAJOR_TYPES = {"header", "section_title", "section_band", "content_box", "exercise",
                   "answer_grid", "table", "items"}

    # Boucle sur les pages et composants
    for page_idx, page in enumerate(pages):
        if page_idx > 0:
            add_page_break(doc)

        components = page.get("components", [])
        prev_type = None
        for component in components:
            comp_type = component.get("type", "").lower()

            # Auto-spacer entre deux composants majeurs consécutifs
            if prev_type in MAJOR_TYPES and comp_type in MAJOR_TYPES:
                add_small_spacer(doc, 2)

            render_component(doc, component, metadata, is_corrige)
            prev_type = comp_type

    # Footer
    if metadata:
        add_footer(
            doc,
            metadata.get("author", "Auteur"),
            metadata.get("doc_type", "Document"),
            metadata.get("number", "01"),
            metadata.get("title", "Titre"),
            metadata.get("level", "Indéterminé")
        )

    return doc


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Génère des documents DOCX style F. Loiseau depuis JSON",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
EXEMPLE :
  python generate_docx.py config.json
  python generate_docx.py config.json --output ./sortie

FORMATS :
  - Génère un DOCX pour chaque version spécifiée dans "versions"
  - Export PDF automatique via LibreOffice ou Word COM
        """
    )
    parser.add_argument("json_file", help="Chemin vers le fichier JSON de configuration")
    parser.add_argument("--output", "-o", help="Répertoire de sortie (défaut: répertoire du JSON)")

    args = parser.parse_args()

    json_path = Path(args.json_file)
    if not json_path.exists():
        print(f"[ERREUR] Fichier JSON introuvable: {json_path}")
        sys.exit(1)

    output_dir = Path(args.output) if args.output else json_path.parent

    print("=" * 70)
    print("Génération de documents DOCX style F. Loiseau")
    print("=" * 70)

    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except json.JSONDecodeError as e:
        print(f"[ERREUR] JSON invalide: {e}")
        sys.exit(1)

    metadata = data.get("metadata", {})
    versions = data.get("versions", ["eleve", "corrige"])

    if not metadata:
        print("[ERREUR] Métadonnées manquantes dans le JSON")
        sys.exit(1)

    title = metadata.get("title", "Document")
    doc_type = metadata.get("doc_type", "Document")
    number = metadata.get("number", "01")

    print(f"\nDocument : {doc_type} {number} – {title}")
    print(f"Niveau : {metadata.get('level', 'Indéterminé')}")
    print(f"Versions : {', '.join(versions)}\n")

    output_dir.mkdir(parents=True, exist_ok=True)

    for version in versions:
        print(f"[1/2] Génération version '{version}'...")
        try:
            doc = create_document_from_json(data, version=version)
            suffix = " – CORRIGÉ" if version == "corrige" else ""
            filename = f"{doc_type}_{number}_{title}{suffix}.docx"
            filename = filename.replace(" ", "_").replace("/", "-")
            output_path = output_dir / filename

            doc.save(str(output_path))
            print(f"  ✓ {output_path.name}")

            # Export PDF
            print(f"[2/2] Export PDF version '{version}'...")
            try:
                result = subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(output_dir),
                    str(output_path)
                ], capture_output=True, text=True, timeout=60)

                if result.returncode == 0:
                    print(f"  ✓ PDF généré")
                else:
                    print(f"  [WARN] LibreOffice: {result.stderr[:100]}")

            except FileNotFoundError:
                print("  [INFO] LibreOffice non trouvé, essai avec Word COM...")
                try:
                    import comtypes.client
                    word = comtypes.client.CreateObject('Word.Application')
                    word.Visible = False
                    doc_word = word.Documents.Open(str(output_path.resolve()))
                    pdf_path = str(output_path.with_suffix('.pdf').resolve())
                    doc_word.SaveAs(pdf_path, FileFormat=17)
                    doc_word.Close()
                    print(f"  ✓ PDF généré via Word COM")
                except Exception as e:
                    print(f"  [SKIP] Pas de convertisseur PDF disponible: {e}")
                finally:
                    try:
                        word.Quit()
                    except:
                        pass

        except Exception as e:
            print(f"  [ERREUR] Génération échouée: {e}")
            import traceback
            traceback.print_exc()

    print("\n" + "=" * 70)
    print("Terminé !")
    print(f"Fichiers générés dans : {output_dir}")
    print("=" * 70)


if __name__ == "__main__":
    main()
